from __future__ import annotations

import io
from datetime import datetime
from typing import Iterable

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from src_code import _draw_4_diagrammi_matplotlib


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8)
    run.font.name = "Arial"


def _add_table(document: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]) -> None:
    headers = list(headers)
    rows = [list(row) for row in rows]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)


def _format_float(value: float, decimals: int = 3) -> str:
    if not np.isfinite(value):
        return "-"
    return f"{value:.{decimals}f}"


def _result_summary_rows(x_m, V, M, theta, v) -> list[list[str]]:
    x_m = np.asarray(x_m, dtype=float)
    V = np.asarray(V, dtype=float)
    M = np.asarray(M, dtype=float)
    theta = np.asarray(theta, dtype=float)
    v = np.asarray(v, dtype=float)

    index_v = int(np.nanargmax(np.abs(V)))
    index_m = int(np.nanargmax(np.abs(M)))
    index_theta = int(np.nanargmax(np.abs(theta)))
    index_def = int(np.nanargmax(np.abs(v)))

    return [
        ["Taglio massimo |V|max", _format_float(abs(V[index_v]) / 1000.0, 3), "kN", _format_float(x_m[index_v], 3)],
        ["Momento massimo |M|max", _format_float(abs(M[index_m]) / 1_000_000.0, 3), "kNm", _format_float(x_m[index_m], 3)],
        ["Rotazione massima |theta|max", _format_float(abs(theta[index_theta]) * 1000.0, 3), "mrad", _format_float(x_m[index_theta], 3)],
        ["Deformata massima |v|max", _format_float(abs(v[index_def]), 3), "mm", _format_float(x_m[index_def], 3)],
    ]


def genera_word_prontuario(
    vincolo: str,
    carico: str,
    parametri_input: dict[str, str],
    risultati_extra: dict[str, str],
    x_m,
    V,
    M,
    theta,
    v,
    L_m: float = 0.0,
) -> bytes:
    """Genera una scheda Word editabile per lo schema statico corrente."""
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Prontuario Strutturale - Scheda di calcolo")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{vincolo} - {carico}").bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    document.add_heading("1. Dati di input", level=2)
    input_rows = [[name, value] for name, value in parametri_input.items()]
    _add_table(document, ["Parametro", "Valore"], input_rows)

    document.add_heading("2. Risultati di sintesi", level=2)
    _add_table(document, ["Grandezza", "Valore", "Unita", "Ascissa x [m]"], _result_summary_rows(x_m, V, M, theta, v))

    if risultati_extra:
        document.add_heading("3. Risultati specifici dello schema", level=2)
        _add_table(document, ["Grandezza", "Valore"], [[name, value] for name, value in risultati_extra.items()])
        diagrams_heading = "4. Diagrammi"
    else:
        diagrams_heading = "3. Diagrammi"

    document.add_heading(diagrams_heading, level=2)
    document.add_paragraph(
        "I diagrammi riportano taglio, momento flettente, rotazione e deformata con le stesse convenzioni della pagina di calcolo."
    )
    diagram_buffer = _draw_4_diagrammi_matplotlib(x_m, V, M, theta, v)
    document.add_picture(diagram_buffer, width=Cm(17.0))

    document.add_heading("Note", level=2)
    document.add_paragraph(
        "La scheda e' un riepilogo automatico del caso selezionato nel prontuario. "
        "Le verifiche progettuali, i coefficienti normativi e la scelta del modello restano a cura del progettista."
    )
    if L_m:
        document.add_paragraph(f"Luce di riferimento usata per la scheda: {L_m:.3f} m.")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
